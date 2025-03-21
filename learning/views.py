import json
import os
import uuid
import requests
import openai
from datetime import timedelta

from django.conf import settings
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.core.exceptions import PermissionDenied
from django.http import JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods, require_POST
from django.contrib import messages
from django.db import transaction

from core import settings
from .forms import DocumentUploadForm, CustomUserCreationForm
from .models import UserProfile, Quiz, Question, Document

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from django.contrib.auth.models import User
from django.contrib.auth.password_validation import validate_password
from django.core.exceptions import ValidationError

def register_view(request):
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        
        try:
            if not username:
                raise ValidationError('Username is required')
            if not password:
                raise ValidationError('Password is required')
            if len(username) < 3:
                raise ValidationError('Username must be at least 3 characters long')
            if User.objects.filter(username=username).exists():
                raise ValidationError('Username already exists')
            
            validate_password(password)
            
            with transaction.atomic():
                # Cleanup any existing entries (if any)
                UserProfile.objects.filter(user__username=username).delete()
                User.objects.filter(username=username).delete()
                
                user = User.objects.create_user(
                    username=username,
                    password=password
                )
                
                # Remove the explicit UserProfile creation if using signals
                # UserProfile.objects.create(user=user)  # <--- Remove this line
                
                messages.success(request, 'Registration successful! Please login.')
                return redirect('login')
                
        except ValidationError as e:
            messages.error(request, str(e.messages[0] if hasattr(e, 'messages') else e))
        except Exception as e:
            messages.error(request, 'Registration failed. Please try again.')
            print(f"Registration error: {str(e)}")
            
            try:
                if 'user' in locals():
                    user.delete()
            except:
                pass
    
    return render(request, 'registration/register.html')
def process_document(document):
    """Extract text from different file formats"""
    text = ""
    try:
        if document.content_type == 'application/pdf':
            pdf_reader = PdfReader(document)
            text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
        elif document.content_type == 'text/plain':
            text = document.read().decode()
        elif document.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = DocxDocument(document)
            text = '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise Exception(f"Error processing document: {str(e)}")
    return text

def generate_summary(text):
    """Generate summary using OpenAI"""
    openai.api_key = os.getenv('OPENAI_API_KEY')
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Generate a comprehensive summary in bullet points with key concepts and important details."},
                {"role": "user", "content": text[:10000]}  # Limit input size
            ],
            temperature=0.3
        )
        return response.choices[0].message['content']
    except Exception as e:
        raise Exception(f"Summary generation failed: {str(e)}")

def generate_quiz(text, user):
    """Generate quiz questions from document content"""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": """Generate 5 quiz questions in JSON format. Each question should have:
                - question: The question text
                - options: 4 possible answers
                - correct_answer: Index of correct option (0-3)
                - explanation: Brief explanation"""},
                {"role": "user", "content": text[:8000]}
            ],
            temperature=0.5
        )
        
        questions_data = json.loads(response.choices[0].message['content'])
        quiz = Quiz.objects.create(
            title=f"Generated Quiz - {user.username}",
            category='GEN',
            difficulty=3,
            user=user
        )
        
        for q in questions_data:
            Question.objects.create(
                quiz=quiz,
                text=q['question'],
                options=q['options'],
                correct_answer=q['correct_answer'],
                explanation=q['explanation']
            )
        
        return quiz
    except Exception as e:
        raise Exception(f"Quiz generation failed: {str(e)}")

@csrf_exempt  # Only for development, use proper CSRF protection in production
@require_POST
def process_voice(request):
    """Handle voice input processing and AI response generation"""
    try:
        # Get audio file from request
        audio_file = request.FILES.get('audio')
        if not audio_file:
            return JsonResponse({'error': 'No audio file provided'}, status=400)

        # Set up OpenAI API
        openai.api_key = os.getenv('OPENAI_API_KEY')
        
        # Step 1: Transcribe audio using Whisper
        with audio_file.open('rb') as audio:
            transcription = openai.Audio.transcribe(
                model="whisper-1",
                file=audio
            )
        
        transcribed_text = transcription['text']

        # Step 2: Get context from active document if provided
        document_context = ""
        if 'document' in request.FILES:
            document = request.FILES['document']
            document_context = process_document(document)  # Your existing document processing function

        # Step 3: Generate AI response using ChatGPT
        messages = [
            {"role": "system", "content": """You are an intelligent learning assistant. 
             Provide clear, concise answers and explain complex concepts simply.
             If relevant, include examples or analogies to aid understanding."""},
            {"role": "user", "content": f"""Context: {document_context[:1000] if document_context else ''}
             Question: {transcribed_text}"""}
        ]

        chat_response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.7,
            max_tokens=500
        )

        ai_response = chat_response.choices[0].message['content']

        # Step 4: Generate audio response using Text-to-Speech
        audio_response = openai.Audio.create(
            model="tts-1",
            voice="alloy",  # You can choose different voices: alloy, echo, fable, onyx, nova, shimmer
            input=ai_response
        )

        # Return both text and audio responses
        return JsonResponse({
            'success': True,
            'transcription': transcribed_text,
            'ai_response': ai_response,
            'audio_url': audio_response.url  # OpenAI provides a temporary URL for the audio
        })

    except Exception as e:
        return JsonResponse({
            'error': f'Error processing voice input: {str(e)}'
        }, status=500)

# Update the existing chat interaction handler to support voice
@csrf_exempt
@require_POST
def handle_chat_interaction(request):
    """Handle chat interactions"""
    print("Received chat request")  # Debug print
    try:
        # Parse JSON data from request body
        try:
            data = json.loads(request.body)
            print("Received data:", data)  # Debug print
            
            message = data.get('message')
            document_id = data.get('documentId')
            
            if not message:
                return JsonResponse({'error': 'No message provided'}, status=400)
                
        except json.JSONDecodeError as e:
            print("JSON decode error:", str(e))  # Debug print
            return JsonResponse({'error': 'Invalid JSON data'}, status=400)
        
        # Get document context if provided
        context = ""
        if document_id and document_id != 'null':
            try:
                document = Document.objects.get(id=document_id)
                context = document.processed_text[:4000]  # Limit context size
            except Document.DoesNotExist:
                print(f"Document not found: {document_id}")  # Debug print
                pass  # Continue without context
            except Exception as e:
                print(f"Error getting document: {str(e)}")  # Debug print
                pass  # Continue without context
        
        # Get AI response using context
        try:
            client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
            
            system_message = """You are a helpful AI tutor. """
            if context:
                system_message += f"""Use the following document context to answer questions:
                {context}
                
                If the question isn't directly related to the document, you can still provide helpful general information."""
            
            # Get main response
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": message}
                ],
                temperature=0.7
            )
            
            ai_response = response.choices[0].message.content
            
            # Generate follow-up question
            followup = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Based on the previous interaction, generate an engaging follow-up question to deepen the user's understanding."},
                    {"role": "user", "content": f"Previous response: {ai_response}"}
                ],
                temperature=0.7
            )
            
            followup_question = followup.choices[0].message.content
            
            return JsonResponse({
                'success': True,
                'response': ai_response,
                'followup_question': followup_question
            })
            
        except Exception as e:
            print(f"OpenAI API error: {str(e)}")  # Debug print
            return JsonResponse({
                'error': f'Error generating response: {str(e)}'
            }, status=500)
    
    except Exception as e:
        print(f"Chat handling error: {str(e)}")  # Debug print
        return JsonResponse({'error': str(e)}, status=500)

# Core Dashboard View
def dashboard(request):
    # Add your dashboard logic here
    context = {
        'user': request.user,
        'stats': {
            'documents_processed': 0,
            'quizzes_taken': 0,
            'average_score': 0
        }
    }
    return render(request, 'dashboard.html', context)

# Document Processing View
@csrf_exempt
@require_POST
def process_document(request):
    """Handle document upload, processing and summarization"""
    try:
        if 'document' not in request.FILES:
            return JsonResponse({'error': 'No document provided'}, status=400)
        
        document = request.FILES['document']
        
        # Validate file type
        allowed_types = [
            'application/pdf',
            'text/plain',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
        
        if document.content_type not in allowed_types:
            return JsonResponse({
                'error': f'Unsupported file type: {document.content_type}'
            }, status=400)
        
        # Extract text from document
        try:
            text = ""
            if document.content_type == 'application/pdf':
                pdf_reader = PdfReader(document)
                text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
            elif document.content_type == 'text/plain':
                text = document.read().decode('utf-8')
            elif document.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = DocxDocument(document)
                text = '\n'.join([para.text for para in doc.paragraphs])
                
            if not text.strip():
                return JsonResponse({
                    'error': 'Could not extract text from document'
                }, status=400)
                
        except Exception as e:
            print(f"Text extraction error: {str(e)}")
            return JsonResponse({
                'error': f'Error extracting text: {str(e)}'
            }, status=400)
        
        # Generate initial summary using OpenAI
        try:
            if not settings.OPENAI_API_KEY:
                return JsonResponse({
                    'error': 'OpenAI API key not configured'
                }, status=500)
                
            client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
            
            # Generate summary
            summary_response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Generate a concise but comprehensive summary of this document, highlighting key points and main ideas."},
                    {"role": "user", "content": text[:4000]}  # Limit input size
                ],
                temperature=0.7
            )
            
            summary = summary_response.choices[0].message.content
            
            # Generate initial questions
            questions_response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Generate 3 engaging questions to ask about this document to test understanding. Format them as a numbered list."},
                    {"role": "user", "content": text[:4000]}
                ],
                temperature=0.7
            )
            
            suggested_questions = questions_response.choices[0].message.content
            
            # Save document if user is authenticated
            if request.user.is_authenticated:
                Document.objects.create(
                    user=request.user,
                    file=document,
                    processed_text=text,
                    summary=summary
                )
            
            return JsonResponse({
                'success': True,
                'name': document.name,
                'summary': summary,
                'suggested_questions': suggested_questions,
                'size': document.size
            })
            
        except Exception as e:
            print(f"OpenAI API error: {str(e)}")
            return JsonResponse({
                'error': f'Error generating summary: {str(e)}'
            }, status=500)
        
    except Exception as e:
        print(f"General error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

def extract_text(document):
    """Extract text from uploaded documentxtract text from uploaded document"""
    text = ""
    try:
        if document.content_type == 'application/pdf':
            pdf_reader = PdfReader(document)
            text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
        elif document.content_type == 'text/plain':
            text = document.read().decode()
        elif document.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = DocxDocument(document)
            text = '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise Exception(f"Error processing document: {str(e)}")
    return text

# Quiz Generation View
@require_POST
def generate_quiz(request):
    try:
        # Add quiz generation logic
        return JsonResponse({'status': 'success', 'quiz_id': 1})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# def register_view(request):
#     if request.method == 'POST':
#         username = request.POST.get('username', '').strip()
#         password = request.POST.get('password', '').strip()
        
#         try:
#             # Validate inputs are not empty
#             if not username:
#                 raise ValidationError('Username is required')
#             if not password:
#                 raise ValidationError('Password is required')
            
#             # Validate username length
#             if len(username) < 3:
#                 raise ValidationError('Username must be at least 3 characters long')
            
#             # Check if username exists
#             if User.objects.filter(username=username).exists():
#                 raise ValidationError('Username already exists')
            
#             # Validate password
#             validate_password(password)
            
#             with transaction.atomic():
#                 # Delete any existing profile for cleanup (if any)
#                 UserProfile.objects.filter(user__username=username).delete()
#                 User.objects.filter(username=username).delete()
                
#                 # Create new user
#                 user = User.objects.create_user(
#                     username=username,
#                     password=password
#                 )
                
#                 # Create new profile
#                 UserProfile.objects.create(user=user)
                
#                 messages.success(request, 'Registration successful! Please login.')
#                 return redirect('login')
                
#         except ValidationError as e:
#             messages.error(request, str(e.messages[0] if hasattr(e, 'messages') else e))
#         except Exception as e:
#             messages.error(request, 'Registration failed. Please try again.')
#             print(f"Registration error: {str(e)}")  # For debugging
            
#             # Cleanup on error
#             try:
#                 if 'user' in locals():
#                     user.delete()
#             except:
#                 pass
    
#     return render(request, 'registration/register.html')

def logout_view(request):
    logout(request)  # This will clear the session
    # Clear any session data
    request.session.flush()
    # Expire the session
    request.session.set_expiry(1)  # Set session to expire immediately
    return redirect('login')

# def login_view(request):
#     if request.method == 'POST':
#         username = request.POST.get('username')
#         password = request.POST.get('password')
#         user = authenticate(request, username=username, password=password)
        
#         if user is not None:
#             login(request, user)
#             # Redirect to the page user was trying to access, or home
#             next_page = request.GET.get('next', 'home')
#             return redirect(next_page)
#         else:
#             messages.error(request, 'Invalid username or password.')
    
#     return render(request, 'registration/login.html')
        # if user is not None:
        #     login(request, user)
        #     return redirect('dashboard')
        # return render(request, 'registration/login.html', {
        #     'error': 'Invalid credentials'
        # })
def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            login(request, user)
            return redirect('dashboard')
        return render(request, 'registration/login.html', {
            'error': 'Invalid credentials'
        })
    
    return render(request, 'registration/login.html')

@login_required
def dashboard(request):
    try:
        context = {
            'user': request.user,
            'stats': {
                'documents_processed': request.user.documents.count(),
                'quizzes_taken': 0,  # Update with your quiz model
                'average_score': request.user.profile.average_score
            }
        }
        return render(request, 'dashboard.html', context)
    except Exception as e:
        # logger.error(f"Dashboard error: {str(e)}")
        print(f"Dashboard error: {str(e)}")
        return render(request, 'error.html', status=500)


@login_required
@require_http_methods(["GET", "POST"])
def documents_view(request):
    """Document management view with form handling"""
    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                document = form.save(commit=False)
                document.user = request.user
                document.save()
                return redirect('documents')
            except Exception as e:
                form.add_error(None, f"Error processing document: {str(e)}")
        return render(request, 'documents.html', {'form': form})
    
    form = DocumentUploadForm()
    return render(request, 'documents.html', {
        'documents': request.user.documents.all(),
        'form': form
    })

@csrf_exempt
@require_POST
def handle_chat(request):
    """Handle document-based chat interactions"""
    try:
        data = json.loads(request.body)
        message = data.get('message', '')
        document_id = data.get('documentId')
        
        # Get document context if provided
        context = ""
        if document_id:
            try:
                document = Document.objects.get(id=document_id)
                context = document.processed_text[:4000]  # Limit context size
            except Document.DoesNotExist:
                pass
        
        # Get AI response using context
        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        
        # Get main response
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": f"""You are a helpful AI tutor. Use the following document context to answer questions.
                 If the question isn't related to the document, you can answer generally.
                 Document context: {context}"""},
                {"role": "user", "content": message}
            ],
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content
        
        # Generate follow-up question
        followup = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Based on the previous interaction, generate an engaging follow-up question to deepen the user's understanding."},
                {"role": "user", "content": f"Previous response: {ai_response}"}
            ],
            temperature=0.7
        )
        
        followup_question = followup.choices[0].message.content
        
        return JsonResponse({
            'success': True,
            'response': ai_response,
            'followup_question': followup_question
        })
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_ai_response(prompt, context=""):
    """Get response from OpenAI's API"""
    openai.api_key = settings.OPENAI_API_KEY
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{
                "role": "system",
                "content": f"You are a helpful tutor. Use this context if relevant: {context}"
            }, {
                "role": "user", 
                "content": prompt
            }],
            temperature=0.7
        )
        return response.choices[0].message['content']
    except Exception as e:
        raise Exception(f"AI request failed: {str(e)}")

def generate_eleven_labs_audio(text, voice_name):
    """Generate audio using Eleven Labs API"""
    try:
        voice_map = settings.ELEVEN_LABS_VOICES
        voice_id = voice_map.get(voice_name, voice_map['Sarah'])
        
        response = requests.post(
            f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}/stream",
            json={"text": text[:1000], "model_id": "eleven_multilingual_v2"},
            headers={"xi-api-key": settings.ELEVEN_LABS_API_KEY}
        )
        
        if response.status_code == 200:
            filename = f"audio/{uuid.uuid4()}.mp3"
            full_path = os.path.join(settings.MEDIA_ROOT, filename)
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
            
            with open(full_path, 'wb') as f:
                f.write(response.content)
            
            return f"{settings.MEDIA_URL}{filename}"
        return None
        
    except Exception as e:
        raise Exception(f"Audio generation failed: {str(e)}")

def analytics_view(request):
    # Add your analytics logic here
    return render(request, 'analytics.html')

@login_required
def chat_view(request):
    return render(request, 'chat.html')

def profile_view(request):
    # Add your profile view logic here
    return render(request, 'profile.html')

def home_view(request):
    return render(request, 'home.html')

